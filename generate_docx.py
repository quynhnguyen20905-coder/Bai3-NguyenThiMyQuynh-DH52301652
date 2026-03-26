from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title_run = title.add_run("BÀI TẬP 03: TÌM HIỂU VỀ RENDER VÀ TRIỂN KHAI TRANG WEB")
title_run.bold = True
title_run.font.size = Pt(16)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Info
doc.add_paragraph("Họ và tên: Nguyễn Thị Mỹ Quỳnh")
doc.add_paragraph("MSSV: DH52301652")
doc.add_paragraph("Lớp: [Tên lớp của bạn]")
doc.add_paragraph("Ngày nộp: 26/03/2026")
doc.add_paragraph("-" * 20)

# Section 1
doc.add_heading("PHẦN 1: TÌM HIỂU VỀ NỀN TẢNG RENDER", level=1)
doc.add_heading("1.1. Render là gì?", level=2)
doc.add_paragraph("Render là một nền tảng đám mây (Cloud Platform) hiện đại, được thiết kế để giúp các nhà phát triển triển khai (deploy), vận hành và mở rộng các ứng dụng web cũng như các trang web tĩnh (static websites) một cách dễ dàng và tự động. Render được xem là một lựa chọn thay thế mạnh mẽ cho Heroku nhờ vào tính đơn giản, hiệu năng cao và chi phí hợp lý (có gói miễn phí cho sinh viên và dự án nhỏ).")

doc.add_heading("1.2. Các tính năng chính của Render", level=2)
features = [
    "Hỗ trợ đa dạng loại hình dịch vụ: Tự động triển khai các trang web tĩnh (HTML/CSS/JS), các dịch vụ web (Node.js, Python, Go, Ruby, v.v.), cơ sở dữ liệu (PostgreSQL, Redis) và các ứng dụng Docker.",
    "Triển khai liên tục (Continuous Deployment - CD): Kết nối trực tiếp với GitHub/GitLab. Mỗi khi bạn đẩy code (push) lên kho lưu trữ, Render sẽ tự động xây dựng (build) và cập nhật trang web mới nhất.",
    "Chứng chỉ SSL miễn phí: Tự động cung cấp và gia hạn chứng chỉ bảo mật HTTPS thông qua Let's Encrypt, giúp trang web an toàn hơn.",
    "Global CDN: Sử dụng mạng lưới phân phối nội dung toàn cầu để tăng tốc độ tải trang cho người dùng ở khắp mọi nơi.",
    "Xem trước Pull Request (Pull Request Previews): Tự động tạo một liên kết xem trước cho mỗi yêu cầu thay đổi (Pull Request) để kiểm tra trước khi gộp vào bản chính.",
    "DDoS Protection: Tích hợp sẵn các lớp bảo vệ chống lại các cuộc tấn công từ chối dịch vụ."
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading("1.3. Ưu điểm của Render", level=2)
advantages = [
    "Dễ sử dụng: Giao diện trực quan, không yêu cầu cấu hình máy chủ phức tạp.",
    "Miễn phí cho dự án nhỏ: Gói Static Site hoàn toàn miễn phí với băng thông lên tới 100GB/tháng.",
    "Hiệu suất cao: Hạ tầng hiện đại giúp ứng dụng chạy ổn định và nhanh chóng."
]
for a in advantages:
    doc.add_paragraph(a, style='List Bullet')

# Section 2
doc.add_heading("PHẦN 2: CÁC BƯỚC TRIỂN KHAI TRANG WEB LÊN RENDER", level=1)
doc.add_heading("Bước 1: Chuẩn bị mã nguồn và Đẩy lên GitHub", level=2)
doc.add_paragraph("Đảm bảo mã nguồn trang web (HTML, CSS và các thư mục hình ảnh) đã được đẩy lên kho lưu trữ GitHub.")
doc.add_paragraph("Link GitHub của project: https://github.com/quynhnguyen20905-coder/Bai3-NguyenThiMyQuynh-DH52301652")

doc.add_heading("Bước 2: Đăng ký tài khoản Render", level=2)
doc.add_paragraph("1. Truy cập vào trang chủ Render.com.\n2. Chọn Sign Up và đăng ký thông qua tài khoản GitHub để đồng bộ hóa kho lưu trữ dễ dàng nhất.")

doc.add_heading("Bước 3: Tạo một Static Site mới trên Render", level=2)
doc.add_paragraph("1. Tại màn hình Dashboard, nhấn nút New + và chọn Static Site.\n2. Kết nối với tài khoản GitHub nếu chưa thực hiện.\n3. Tìm và chọn kho lưu trữ có tên: Bai3-NguyenThiMyQuynh-DH52301652.")

doc.add_heading("Bước 4: Cấu hình thông số triển khai", level=2)
doc.add_paragraph("Name: Đặt tên cho dự án (ví dụ: my-quynh-dh52301652).\nBranch: Chọn nhánh main.\nRoot Directory: Để trống.\nBuild Command: Để trống.\nPublish Directory: Điền dấu chấm ( . )")

doc.add_heading("Bước 5: Triển khai và Kiểm tra", level=2)
doc.add_paragraph("1. Nhấn Create Static Site.\n2. Chờ đợi Render thực hiện quá trình Build và Deploy (thường mất khoảng 1-2 phút).\n3. Sau khi trạng thái chuyển sang Live, bạn sẽ nhận được một đường link (ví dụ: https://my-quynh-dh52301652.onrender.com).")

# Section 3
doc.add_heading("PHẦN 3: KẾT QUẢ VÀ ĐƯỜNG LINK KIỂM TRA", level=1)
doc.add_paragraph("Đường link GitHub Repo: https://github.com/quynhnguyen20905-coder/Bai3-NguyenThiMyQuynh-DH52301652")
doc.add_paragraph("Đường link Live Web (trên Render): [Điền link của bạn tại đây]")

doc.save("Baitap03_Timhieu_Render.docx")
